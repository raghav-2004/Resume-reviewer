"""
Resume Parser Service
Extracts text from PDF and DOCX resume files
"""

import io
import re
from typing import Optional


def clean_text(text: str) -> str:
    """Remove excessive whitespace and clean up extracted text."""
    # Replace multiple newlines with double newline
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Replace multiple spaces with single space
    text = re.sub(r' {2,}', ' ', text)
    # Remove non-printable characters
    text = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]', '', text)
    return text.strip()


def parse_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF using pdfplumber.
    Falls back to PyPDF2 if pdfplumber fails.
    """
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_text = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
            full_text = "\n\n".join(pages_text)
            return clean_text(full_text)
    except Exception as e:
        # Fallback to pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
            full_text = "\n\n".join(pages_text)
            return clean_text(full_text)
        except Exception as e2:
            raise ValueError(f"Could not parse PDF: {str(e)} | Fallback error: {str(e2)}")


def parse_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX using python-docx.
    Preserves paragraph structure.
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return clean_text("\n".join(paragraphs))
    except Exception as e:
        raise ValueError(f"Could not parse DOCX: {str(e)}")


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """
    Auto-detect file type and parse resume.
    Returns extracted text content.
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif filename_lower.endswith(".docx") or filename_lower.endswith(".doc"):
        return parse_docx(file_bytes)
    elif filename_lower.endswith(".txt"):
        return clean_text(file_bytes.decode("utf-8", errors="replace"))
    else:
        raise ValueError(f"Unsupported file type: {filename}. Please upload PDF, DOCX, or TXT.")
